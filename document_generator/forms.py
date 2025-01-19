from django import forms
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document

class CamposDinamicosForm(forms.Form):
    processo_administrativo = forms.CharField(label="Número do Processo Administrativo", max_length=50)
    objeto = forms.CharField(label="Descrição do Objeto", max_length=200)
    tipo_vigencia = forms.ChoiceField(
        label="Prazo de Vigência",
        choices=[
            ("a", "Entrega única (Art. 105)"),
            ("b", "Entrega contínua (Art. 106)"),
            ("c", "Emergencial (Art. 75, VIII)"),
        ],
    )
    num_criterios = forms.IntegerField(label="Quantidade de Critérios de Sustentabilidade", min_value=1, initial=1)
    luxury_item = forms.ChoiceField(
        label="O objeto é um bem de luxo?",
        choices=[("no", "Não"), ("yes", "Sim")],
    )
    justificativa = forms.CharField(
        label="Justificativa da Contratação",
        widget=forms.Textarea,
        required=False,
        help_text="Deixe em branco para usar a justificativa padrão."
    )
    plano_contratacoes = forms.ChoiceField(
        label="A contratação está prevista no PCA?",
        choices=[("sim", "Sim"), ("nao", "Não")],
    )

class ItensForm(forms.Form):
    num_itens = forms.IntegerField(label="Quantidade de Itens", min_value=1, initial=1)
    num_lotes = forms.IntegerField(label="Quantidade de Lotes", min_value=1, initial=1)

def preencher_documento(request):
    if request.method == "POST":
        # Coleta os dados do formulário
        form_dinamico = CamposDinamicosForm(request.POST)
        form_itens = ItensForm(request.POST)

        if form_dinamico.is_valid() and form_itens.is_valid():
            # Dados dos formulários
            dados_dinamicos = form_dinamico.cleaned_data
            num_itens = form_itens.cleaned_data["num_itens"]
            num_lotes = form_itens.cleaned_data["num_lotes"]

            # Carrega o template do documento
            doc = Document("template.docx")

            # Preenche os placeholders gerais
            preencher_placeholders(doc, dados_dinamicos)

            # Adiciona itens e lotes na tabela
            adicionar_itens_e_lotes(doc, num_itens, num_lotes)

            # Valida os placeholders
            validate_placeholders(doc)

            # Retorna o documento gerado
            response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            response["Content-Disposition"] = 'attachment; filename="projeto_basico.docx"'
            doc.save(response)
            return response

    else:
        form_dinamico = CamposDinamicosForm()
        form_itens = ItensForm()

    return render(request, "formulario.html", {"form_dinamico": form_dinamico, "form_itens": form_itens})


def preencher_placeholders(doc, dados):
    """Substitui os placeholders no documento com os dados fornecidos."""
    placeholders = {
        "<<processo_administrativo>>": dados.get("processo_administrativo"),
        "<<objeto>>": dados.get("objeto"),
        "<<tipo_vigencia>>": dados.get("tipo_vigencia"),
        "<<justificativa>>": dados.get("justificativa", "A Fundamentação da Contratação encontra-se no apêndice."),
        "<<plano_contratacoes>>": "Sim" if dados.get("plano_contratacoes") == "sim" else "Não",
    }
    for p in doc.paragraphs:
        for placeholder, valor in placeholders.items():
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, valor)


def adicionar_itens_e_lotes(doc, num_itens, num_lotes):
    """Adiciona itens e lotes dinamicamente na tabela do documento."""
    table = doc.tables[0]  # Supondo que a tabela é a primeira
    for lote in range(1, num_lotes + 1):
        table.add_row().cells[0].text = f"LOTE {lote}"
        for item in range(1, num_itens + 1):
            row = table.add_row().cells
            row[0].text = str(item)
            row[1].text = f"Descrição do Item {item} do Lote {lote}"
            row[2].text = "Unidade"
            row[3].text = "Quantidade"
            row[4].text = "Preço Unitário"
            row[5].text = "Preço Total"


def validate_placeholders(doc):
    """Valida se todos os placeholders foram preenchidos."""
    for p in doc.paragraphs:
        if "<<" in p.text and ">>" in p.text:
            raise ValueError(f"Placeholder não preenchido: {p.text}")
